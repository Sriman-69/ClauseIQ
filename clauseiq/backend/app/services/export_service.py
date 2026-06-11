import os
import uuid
import json as pyjson
import asyncio
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document as DocxDocument
from app.repositories.document_repository import DocumentRepository
from app.services.summary_service import SummaryService
from app.services.checklist_service import ChecklistService
from app.services.risk_service import RiskService

class ExportService:
    def __init__(self, db):
        self.db = db
        self.document_repo = DocumentRepository(db)
        self.summary_service = SummaryService(db)
        self.checklist_service = ChecklistService(db)
        self.risk_service = RiskService(db)

    async def export_report(self, document_id: str, user_id: str, export_format: str = 'pdf') -> dict:
        document = self.document_repo.get_by_id(document_id, user_id=user_id)
        if not document:
            raise ValueError("Document not found")

        # Reads directly from cache because of Snapshot implementation!
        summary, checklist, risks = await asyncio.gather(
            self.summary_service.generate_summary(document_id, user_id=user_id),
            self.checklist_service.generate_checklist(document_id, user_id=user_id),
            self.risk_service.analyze_risks(document_id, user_id=user_id)
        )

        EXPORT_DIR = "exports"
        os.makedirs(EXPORT_DIR, exist_ok=True)
        
        if export_format == 'json':
            report_filename = f"report_{uuid.uuid4().hex}.json"
            report_path = os.path.join(EXPORT_DIR, report_filename)
            with open(report_path, "w") as f:
                pyjson.dump({"summary": summary, "checklist": checklist, "risks": risks}, f, indent=4)
            return {"download_url": f"/api/v1/export/download/{report_filename}"}
            
        elif export_format == 'docx':
            report_filename = f"report_{uuid.uuid4().hex}.docx"
            report_path = os.path.join(EXPORT_DIR, report_filename)
            doc = DocxDocument()
            doc.add_heading(f"ClauseIQ Analysis Report: {document.filename}", 0)
            doc.add_paragraph("ClauseIQ provides AI-assisted research support only and does not constitute legal advice.")
            
            doc.add_heading("1. Document Summary", level=1)
            doc.add_paragraph(f"Executive Summary: {summary.get('executive_summary', '')}")
            
            doc.add_heading("2. Compliance Checklist", level=1)
            for item in checklist:
                doc.add_paragraph(f"• {item.get('title')} ({item.get('status')}): {item.get('explanation')}")
                
            doc.add_heading("3. Risk Analysis", level=1)
            for r in risks.get('high_risks', []):
                doc.add_paragraph(f"[HIGH] {r.get('risk')}: {r.get('reason')}")
            for r in risks.get('medium_risks', []):
                doc.add_paragraph(f"[MEDIUM] {r.get('risk')}: {r.get('reason')}")
                
            doc.save(report_path)
            return {"download_url": f"/api/v1/export/download/{report_filename}"}

        else: # PDF
            report_filename = f"report_{uuid.uuid4().hex}.pdf"
            report_path = os.path.join(EXPORT_DIR, report_filename)

            doc = SimpleDocTemplate(report_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            title_style = styles['Heading1']
            h2_style = styles['Heading2']
            normal_style = styles['Normal']

            story.append(Paragraph("ClauseIQ provides AI-assisted research support only.", normal_style))
            story.append(Spacer(1, 12))
            story.append(Paragraph(f"ClauseIQ Analysis Report: {document.filename}", title_style))
            story.append(Spacer(1, 12))

            story.append(Paragraph("1. Document Summary", h2_style))
            story.append(Paragraph(f"<b>Executive Summary:</b> {summary.get('executive_summary', '')}", normal_style))
            story.append(Spacer(1, 12))

            story.append(Paragraph("2. Compliance Checklist", h2_style))
            for item in checklist:
                text = f"• <b>{item.get('title')}</b> ({item.get('status')}): {item.get('explanation')}"
                story.append(Paragraph(text, normal_style))
                story.append(Spacer(1, 6))
            story.append(Spacer(1, 12))

            story.append(Paragraph("3. Risk Analysis", h2_style))
            for risk in risks.get('high_risks', []):
                text = f"• <b>[HIGH] {risk.get('risk')}</b>: {risk.get('reason')}"
                story.append(Paragraph(text, normal_style))
                story.append(Spacer(1, 6))

            doc.build(story)
            return {"download_url": f"/api/v1/export/download/{report_filename}"}
